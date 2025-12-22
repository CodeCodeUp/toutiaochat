"""
DOCX 文件生成工具

用于将 Markdown 格式的文章内容（含图片）转换为 Word 文档格式
使用 pypandoc 处理 Markdown 格式转换
"""

import os
import re
import tempfile
import uuid
from pathlib import Path

import pypandoc
import structlog
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

logger = structlog.get_logger()


class DocxGenerator:
    """DOCX 文档生成器（支持 Markdown 格式和图片嵌入）"""

    def __init__(self):
        self.temp_dir = Path(tempfile.gettempdir()) / "toutiao_articles"
        self.temp_dir.mkdir(exist_ok=True)

    def _set_chinese_font(self, run, font_name: str = "宋体"):
        """设置中文字体"""
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    def _add_image_to_doc(
        self,
        doc: Document,
        image_path: str,
        width_inches: float = 5.5,
        caption: str = None,
    ) -> bool:
        """
        添加图片到文档

        Args:
            doc: 文档对象
            image_path: 图片本地路径
            width_inches: 图片宽度（英寸）
            caption: 图片说明（可选）

        Returns:
            bool: 是否成功添加
        """
        try:
            if not os.path.exists(image_path):
                logger.warning("image_not_found", path=image_path)
                return False

            # 添加图片
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(image_path, width=Inches(width_inches))

            # 添加图片说明
            if caption:
                caption_para = doc.add_paragraph()
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption_para.add_run(caption)
                caption_run.font.size = Pt(9)
                caption_run.font.italic = True
                self._set_chinese_font(caption_run, "宋体")

            return True
        except Exception as e:
            logger.error("add_image_failed", path=image_path, error=str(e))
            return False

    def _insert_image_after_element(
        self,
        doc: Document,
        element,
        image_path: str,
        width_inches: float = 5.0,
    ) -> bool:
        """
        在指定元素后插入图片

        Args:
            doc: 文档对象
            element: 要在其后插入的元素
            image_path: 图片路径
            width_inches: 图片宽度
        """
        try:
            if not os.path.exists(image_path):
                logger.warning("image_not_found", path=image_path)
                return False

            # 创建新段落用于图片
            new_para = doc.add_paragraph()
            new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = new_para.add_run()
            run.add_picture(image_path, width=Inches(width_inches))

            # 将新段落移动到目标元素之后
            element._element.addnext(new_para._element)

            return True
        except Exception as e:
            logger.error("insert_image_failed", path=image_path, error=str(e))
            return False

    def _organize_images_by_position(self, images: list) -> dict:
        """
        按位置组织图片

        Returns:
            dict: {
                "cover": [img1, ...],
                "after_paragraph": {1: [img], 3: [img], ...},
                "end": [img1, ...]
            }
        """
        organized = {
            "cover": [],
            "after_paragraph": {},
            "end": [],
        }

        for img in images:
            position = img.get("position", "end")
            path = img.get("path", "")

            if not path or not os.path.exists(path):
                continue

            if position == "cover":
                organized["cover"].append(img)
            elif position.startswith("after_paragraph:"):
                try:
                    para_num = int(position.split(":")[1])
                    if para_num not in organized["after_paragraph"]:
                        organized["after_paragraph"][para_num] = []
                    organized["after_paragraph"][para_num].append(img)
                except:
                    organized["end"].append(img)
            else:
                organized["end"].append(img)

        return organized

    def _count_content_paragraphs(self, content: str) -> int:
        """统计 Markdown 内容中的段落数（排除标题和空行）"""
        lines = content.split('\n')
        count = 0
        for line in lines:
            line = line.strip()
            # 跳过空行和标题
            if not line or line.startswith('#'):
                continue
            count += 1
        return count

    def _convert_md_to_docx(self, title: str, content: str) -> str:
        """
        使用 pypandoc 将 Markdown 转换为 DOCX

        Args:
            title: 文章标题
            content: Markdown 格式的正文

        Returns:
            str: 临时 DOCX 文件路径
        """
        # 预处理：在每个空行中插入占位符，防止被 Markdown 合并
        # 匹配后面紧跟换行的换行符，插入不间断空格
        # \n\n → \n\u00A0\n, \n\n\n → \n\u00A0\n\u00A0\n
        processed_content = re.sub(r'\n(?=\n)', '\n\u00A0', content)

        # 组合标题和内容
        full_md = f"# {title}\n\n{processed_content}"

        # 生成临时文件路径
        temp_path = str(self.temp_dir / f"temp_{uuid.uuid4().hex[:8]}.docx")

        try:
            # 使用 hard_line_breaks 扩展，让所有换行都被保留
            pypandoc.convert_text(
                full_md,
                'docx',
                format='markdown+hard_line_breaks',
                outputfile=temp_path,
            )
            return temp_path
        except Exception as e:
            logger.error("pypandoc_convert_failed", error=str(e))
            raise

    def create_article_docx(
        self,
        title: str,
        content: str,
        images: list = None,
        article_id: str = None,
    ) -> str:
        """
        创建文章 DOCX 文件（含图片）

        Args:
            title: 文章标题
            content: Markdown 格式的文章正文
            images: 图片列表 [{"path": str, "position": str, "prompt": str}, ...]
            article_id: 文章ID (用于文件名)

        Returns:
            str: DOCX 文件的完整路径
        """
        images = images or []

        # 1. 使用 pypandoc 转换 Markdown 为 DOCX
        temp_docx_path = self._convert_md_to_docx(title, content)

        # 2. 打开生成的 DOCX
        doc = Document(temp_docx_path)

        # 3. 组织图片
        organized_images = self._organize_images_by_position(images)

        # 4. 获取所有段落（跳过标题）
        paragraphs = list(doc.paragraphs)

        # 找到第一个非标题段落的索引
        first_content_idx = 0
        for i, para in enumerate(paragraphs):
            if para.style.name.startswith('Heading'):
                first_content_idx = i + 1
            else:
                break

        # 5. 在标题后插入封面图
        if organized_images["cover"] and first_content_idx > 0:
            # 找到标题段落
            title_para = paragraphs[first_content_idx - 1]
            for img in reversed(organized_images["cover"]):
                self._insert_image_after_element(doc, title_para, img["path"], 5.5)

        # 6. 在指定段落后插入图片
        # 统计正文段落（非标题）
        content_para_count = 0
        for i, para in enumerate(paragraphs):
            if i < first_content_idx:
                continue
            if para.style.name.startswith('Heading'):
                continue

            content_para_count += 1

            # 检查是否需要在此段落后插入图片
            if content_para_count in organized_images["after_paragraph"]:
                for img in reversed(organized_images["after_paragraph"][content_para_count]):
                    self._insert_image_after_element(doc, para, img["path"], 5.0)

        # 7. 在文档末尾添加结尾图片
        for img in organized_images["end"]:
            self._add_image_to_doc(doc, img["path"], width_inches=5.0)

        # 8. 应用中文字体样式
        for para in doc.paragraphs:
            for run in para.runs:
                if not para.style.name.startswith('Heading'):
                    run.font.size = Pt(12)
                    self._set_chinese_font(run, "宋体")
                else:
                    self._set_chinese_font(run, "黑体")

        # 9. 生成最终文件名
        if article_id:
            filename = f"article_{article_id}.docx"
        else:
            filename = f"article_{uuid.uuid4().hex[:8]}.docx"

        # 10. 保存最终文件
        file_path = self.temp_dir / filename
        doc.save(str(file_path))

        # 清理临时文件
        try:
            os.remove(temp_docx_path)
        except:
            pass

        logger.info(
            "docx_created",
            path=str(file_path),
            title=title[:30],
            image_count=len(images),
        )

        return str(file_path)

    def create_preview_docx(
        self,
        title: str,
        content: str,
        images: list = None,
        article_id: str = None,
    ) -> str:
        """
        创建预览用的 DOCX 文件（与发布版相同）

        这是 create_article_docx 的别名，用于语义清晰
        """
        return self.create_article_docx(title, content, images, article_id)

    def get_docx_path(self, article_id: str) -> str | None:
        """获取已生成的 DOCX 文件路径"""
        file_path = self.temp_dir / f"article_{article_id}.docx"
        if file_path.exists():
            return str(file_path)
        return None

    def cleanup_old_files(self, max_age_hours: int = 24):
        """
        清理旧的临时文件

        Args:
            max_age_hours: 文件保留时间(小时)
        """
        import time

        if not self.temp_dir.exists():
            return

        current_time = time.time()
        max_age_seconds = max_age_hours * 3600
        cleaned = 0

        for file_path in self.temp_dir.glob("*.docx"):
            try:
                file_age = current_time - os.path.getmtime(file_path)
                if file_age > max_age_seconds:
                    file_path.unlink()
                    cleaned += 1
            except Exception:
                pass

        if cleaned > 0:
            logger.info("docx_cleanup", cleaned_count=cleaned)


# 全局实例
docx_generator = DocxGenerator()
